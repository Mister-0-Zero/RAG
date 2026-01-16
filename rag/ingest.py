"""
Provides functions for ingesting and processing documents from a directory.
"""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
import fitz
from pydantic import BaseModel

from rag.config import RAGConfig

from support_function.detect_function import detect_category

log = logging.getLogger(__name__)

def normalize_text(text: str) -> str:
    """Normalizes text by standardizing line endings and removing trailing whitespace."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = [line.rstrip() for line in text.split("\n")]

    normalized_lines: list[str] = []

    for line in lines:
        if line != "":
            normalized_lines.append(line)

    return "\n".join(normalized_lines).strip()

class RawDocument(BaseModel):
    """A metadata class representing a raw document."""
    id: str
    path: Path
    text: str
    source: str
    doc_type: str
    category: str | None = None
    allowed_roles: str


def read_txt(path: Path) -> str:
    """Reads text content from a .txt file."""
    with path.open(encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx(path: Path) -> str:
    """Reads text content from a .docx file."""
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    return "\n".join(texts)


def read_pdf(path: Path) -> str:
    """Reads text content from a .pdf file."""
    texts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            texts.append(page.get_text())
    return "\n".join(texts)


def ingest_directory(path_dir: Path, cfg: RAGConfig) -> list[RawDocument]:
    log.info(f"Starting ingestion from directory: {path_dir}...", extra={'log_type': 'INFO'})

    documents: list[RawDocument] = []
    supported_files = 0
    unsupported_files = 0

    acl_rules = None
    if cfg.acl:
        from rag.acl_rules import ACLRules
        acl_rules = ACLRules.load(cfg.acl_rules_path)

    for file in path_dir.rglob("*"):
        if not file.is_file():
            continue

        suffix = file.suffix.lower()

        if suffix == ".txt":
            raw_text = read_txt(file)
            doc_type = "txt"
        elif suffix == ".docx":
            raw_text = read_docx(file)
            doc_type = "docx"
        elif suffix == ".pdf":
            raw_text = read_pdf(file)
            doc_type = "pdf"
        else:
            unsupported_files += 1
            continue

        text = normalize_text(raw_text)
        category = detect_category(file)
        supported_files += 1

        if cfg.acl and acl_rules:
            relative_path = file.relative_to(path_dir).as_posix()
            roles = acl_rules.resolve_roles(relative_path)

            if roles is not None:
                allowed_roles = "|".join(roles)
            else:
                allowed_roles = "*" if cfg.default_allow else ""
        else:
            allowed_roles = "*"

        log.debug(
            "Ingesting file: %s, assigned roles: %s",
            file.name,
            allowed_roles,
        )

        documents.append(
            RawDocument(
                id=file.stem,
                path=file,
                text=text,
                source=file.name,
                doc_type=doc_type,
                category=category,
                allowed_roles=allowed_roles,
            )
        )

    log.info(
        f"Finished ingestion. Ingested {supported_files} supported files. "
        f"Skipped {unsupported_files} unsupported files.",
        extra={'log_type': 'INFO'}
    )
    return documents



def ingest_all(cfg: RAGConfig | None = None) -> list[RawDocument]:
    """A wrapper function to start the ingestion process using a `RAGConfig`."""
    cfg = cfg or RAGConfig()
    return ingest_directory(cfg.data_raw, cfg)